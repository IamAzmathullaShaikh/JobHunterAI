import logging
import os
from typing import Any, Dict

from jinja2 import Template

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class TemplateEngine:
    """
    Export Engine for 10 Professional Resume Templates.
    Supports pixel-perfect HTML rendering and DOCX export.
    """

    BASE_CSS = """
    :root {
        --accent-color: {{ config.accent_color | default('#4f46e5') }};
        --font-family: {{ config.font_family | default("'Inter', sans-serif") }};
        --font-size: {{ config.font_size | default('13px') }};
        --line-height: {{ config.line_height | default('1.5') }};
        --margin: {{ config.margin | default('40px') }};
    }
    body { font-family: var(--font-family); color: #1f2937; line-height: var(--line-height); margin: 0; padding: var(--margin); font-size: var(--font-size); }
    h1 { font-size: 2em; margin: 0 0 4px 0; color: #111827; }
    h2 { font-size: 1.2em; text-transform: uppercase; letter-spacing: 0.1em; border-bottom: 2px solid var(--accent-color); padding-bottom: 4px; margin: 24px 0 12px 0; color: #374151; }
    h3 { font-size: 1.1em; margin: 0; color: #111827; }
    .header-sub { font-size: 0.9em; color: #6b7280; margin-bottom: 20px; }
    .section-item { margin-bottom: 16px; page-break-inside: avoid; }
    .item-header { display: flex; justify-content: space-between; font-weight: bold; }
    .item-sub { display: flex; justify-content: space-between; font-style: italic; font-size: 0.95em; color: #4b5563; }
    .bullets { margin: 8px 0 0 20px; padding: 0; font-size: 1em; }
    .bullets li { margin-bottom: 4px; }
    .skills-list { display: flex; flex-wrap: wrap; gap: 8px; font-size: 0.9em; margin-top: 8px; }
    .skill-tag { background: #f3f4f6; padding: 2px 8px; border-radius: 4px; border: 1px solid #e5e7eb; }
    @media print {
        body { padding: 0; }
        .no-print { display: none; }
    }
    """

    TEMPLATES = {
        "classic_ats": """
        <html>
        <head><style>{{ css }}</style></head>
        <body>
            <div style="text-align: center;">
                <h1 style="margin: 0; padding: 0;">{{ header.name }}</h1>
                <div class="header-sub">
                    {{ header.email }} | {{ header.phone }} | {{ header.location }}<br>
                    {% if header.linkedin %}<a href="{{ header.linkedin }}">LinkedIn</a>{% endif %}
                    {% if header.github %} | <a href="{{ header.github }}">GitHub</a>{% endif %}
                </div>
            </div>

            {% if summary %}
            <h2>Professional Summary</h2>
            <p>{{ summary }}</p>
            {% endif %}

            {% if work_history %}
            <h2>Professional Experience</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div class="item-header"><span>{{ job.company }}</span><span>{{ job.start_date }} - {{ job.end_date }}</span></div>
                <div class="item-sub"><span>{{ job.title }}</span><span>{{ job.location }}</span></div>
                <ul class="bullets">
                    {% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}
                </ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if skills %}
            <h2>Core Competencies</h2>
            <div class="skills-list">
                {% for s in skills %}<span class="skill-tag">{{ s }}</span>{% endfor %}
            </div>
            {% endif %}

            {% if projects %}
            <h2>Key Projects</h2>
            {% for p in projects %}
            <div class="section-item">
                <div class="item-header"><span>{{ p.name }}</span><span>{{ p.date }}</span></div>
                <div class="item-sub"><span>{{ p.role }}</span></div>
                <ul class="bullets">
                    {% for b in p.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}
                </ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if education %}
            <h2>Education</h2>
            {% for edu in education %}
            <div class="section-item">
                <div class="item-header"><span>{{ edu.school }}</span><span>{{ edu.date }}</span></div>
                <div class="item-sub"><span>{{ edu.degree }}</span><span>{{ edu.location }}</span></div>
            </div>
            {% endfor %}
            {% endif %}

            {% if certifications %}
            <h2>Certifications</h2>
            <ul class="bullets">
                {% for c in certifications %}<li><b>{{ c.name }}</b>, {{ c.issuer }} ({{ c.date }})</li>{% endfor %}
            </ul>
            {% endif %}
        </body>
        </html>
        """,
        "modern_minimal": """
        <html>
        <head><style>
            {{ css }}
            h1 { font-size: 32px; border-left: 8px solid var(--accent-color); padding-left: 20px; margin-bottom: 20px; }
            h2 { border-bottom: none; border-left: 4px solid var(--accent-color); padding-left: 12px; font-size: 16px; margin-top: 30px; }
            .section-item { margin-left: 16px; border-left: 1px solid #eee; padding-left: 20px; }
            .skill-tag { border-radius: 99px; background: white; border: 1px solid var(--accent-color); color: var(--accent-color); }
        </style></head>
        <body>
            <h1>{{ header.name }}</h1>
            <div class="header-sub" style="font-weight: bold; color: var(--accent-color);">{{ header.title }}</div>
            <div class="header-sub">{{ header.email }} • {{ header.phone }} • {{ header.location }}</div>

            {% if summary %}<h2>Background</h2><p>{{ summary }}</p>{% endif %}

            {% if work_history %}
            <h2>Experience</h2>
            {% for job in work_history %}
            <div class="section-item">
                <h3 style="color: var(--accent-color);">{{ job.title }}</h3>
                <div class="item-sub" style="font-weight: bold; font-style: normal;">{{ job.company }} <span style="color: #ccc;">|</span> {{ job.start_date }} - {{ job.end_date }}</div>
                <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if projects %}
            <h2>Projects</h2>
            {% for p in projects %}
            <div class="section-item">
                <h3>{{ p.name }}</h3>
                <div class="item-sub" style="font-style: italic;">{{ p.role }} | {{ p.date }}</div>
                <ul class="bullets">{% for b in p.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if education %}
            <h2>Education</h2>
            {% for edu in education %}
            <div class="section-item">
                <h3>{{ edu.school }}</h3>
                <p>{{ edu.degree }} | {{ edu.date }}</p>
            </div>
            {% endfor %}
            {% endif %}

            {% if skills %}
            <h2>Expertise</h2>
            <div class="skills-list">{% for s in skills %}<span class="skill-tag">{{ s }}</span>{% endfor %}</div>
            {% endif %}

            {% if certifications %}
            <h2>Certifications</h2>
            <ul class="bullets">
                {% for c in certifications %}<li>{{ c.name }}, {{ c.issuer }} ({{ c.date }})</li>{% endfor %}
            </ul>
            {% endif %}
        </body>
        </html>
        """,
        "executive_elegant": """
        <html>
        <head><style>
            {{ css }}
            body { font-family: 'Georgia', serif; background-color: #fffaf5; }
            h1 { font-family: 'Inter', sans-serif; font-weight: 900; text-transform: uppercase; border-bottom: 4px solid #111827; letter-spacing: -1px; }
            h2 { border-bottom-color: #111827; font-family: 'Inter', sans-serif; font-weight: 800; }
            .accent { color: var(--accent-color); }
        </style></head>
        <body>
            <div style="display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 40px;">
                <h1 style="border:none; margin:0;">{{ header.name }}</h1>
                <div class="header-sub" style="text-align: right; margin:0;">{{ header.location }}<br>{{ header.email }}<br>{{ header.phone }}</div>
            </div>
            <div style="background: var(--accent-color); height: 4px; width: 100%; margin-bottom: 40px;"></div>

            {% if summary %}<p style="font-style: italic; font-size: 1.1em; line-height: 1.6; border-left: 4px solid #eee; padding-left: 20px;">{{ summary }}</p>{% endif %}

            {% if work_history %}
            <h2>Professional Experience</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div class="item-header"><span class="accent" style="font-size: 1.1em;">{{ job.company | upper }}</span><span>{{ job.start_date }} - {{ job.end_date }}</span></div>
                <div class="item-sub" style="font-weight: bold;">{{ job.title }}</div>
                <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if projects %}
            <h2>Key Initiatives</h2>
            {% for p in projects %}
            <div class="section-item">
                <div class="item-header"><span class="accent">{{ p.name }}</span><span>{{ p.date }}</span></div>
                <ul class="bullets">{% for b in p.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if education %}
            <h2>Education</h2>
            {% for edu in education %}
            <div class="section-item">
                <div class="item-header"><span>{{ edu.school }}</span><span>{{ edu.date }}</span></div>
                <p>{{ edu.degree }}</p>
            </div>
            {% endfor %}
            {% endif %}

            {% if skills %}
            <h2>Core Competencies</h2>
            <p>{{ skills | join(' • ') }}</p>
            {% endif %}
        </body>
        </html>
        """,
        "tech_clean": """
        <html>
        <head><style>
            {{ css }}
            body { font-family: 'Roboto Mono', monospace; font-size: 12px; color: #00ff41; background: #0d0208; }
            h1 { color: var(--accent-color); border: 1px solid var(--accent-color); display: inline-block; padding: 10px 20px; }
            h2 { background: var(--accent-color); color: #0d0208; padding: 4px 8px; border-bottom: none; font-weight: bold; }
            .skill-tag { border-radius: 0; border: 1px solid var(--accent-color); background: none; color: #00ff41; }
            a { color: var(--accent-color); }
            .item-sub { color: #888; }
            .bullets { list-style: none; }
            .bullets li::before { content: "> "; color: var(--accent-color); }
        </style></head>
        <body>
            <div style="margin-bottom: 40px;">
                <h1>{{ header.name }}</h1>
                <p style="margin-top: 20px;">SYSTEM_ROLE: {{ header.title }}</p>
                <p>CONTACT: {{ header.email }} | {{ header.github or 'N/A' }} | {{ header.website or 'N/A' }}</p>
            </div>

            {% if summary %}<h2># profile_summary</h2><p>{{ summary }}</p>{% endif %}

            {% if work_history %}
            <h2># experience_logs</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div class="item-header"><span style="color: var(--accent-color);">[{{ job.company }}]</span><span>{{ job.start_date }}..{{ job.end_date }}</span></div>
                <div class="item-sub">{{ job.title }} @ {{ job.location }}</div>
                <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if projects %}
            <h2># technical_projects</h2>
            {% for p in projects %}
            <div class="section-item">
                <div class="item-header"><span style="color: var(--accent-color);">{{ p.name }}</span><span>{{ p.date }}</span></div>
                <ul class="bullets">{% for b in p.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if skills %}
            <h2># technical_stack</h2>
            <div class="skills-list">{% for s in skills %}<span class="skill-tag">{{ s }}</span>{% endfor %}</div>
            {% endif %}
        </body>
        </html>
        """,
        "compact": """
        <html>
        <head><style>
            {{ css }}
            body { padding: 20px; font-size: 11px; }
            h1 { font-size: 22px; text-align: center; }
            h2 { margin: 8px 0 4px 0; border-bottom: 1px solid #ddd; }
            .section-item { margin-bottom: 4px; }
            .bullets { margin-top: 2px; }
            .header-sub { text-align: center; margin-bottom: 10px; }
        </style></head>
        <body>
            <h1>{{ header.name }}</h1>
            <div class="header-sub">{{ header.email }} | {{ header.phone }} | {{ header.location }}</div>

            {% if summary %}<h2>Summary</h2><p>{{ summary }}</p>{% endif %}

            {% if work_history %}
            <h2>Experience</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div class="item-header"><span><b>{{ job.company }}</b> - {{ job.title }}</span><span>{{ job.start_date }} - {{ job.end_date }}</span></div>
                <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if education %}
            <h2>Education</h2>
            {% for edu in education %}
            <div class="section-item"><b>{{ edu.school }}</b>, {{ edu.degree }} ({{ edu.date }})</div>
            {% endfor %}
            {% endif %}

            {% if skills %}
            <h2>Skills</h2>
            <p>{{ skills | join(', ') }}</p>
            {% endif %}
        </body>
        </html>
        """,
        "two_column": """
        <html>
        <head><style>
            {{ css }}
            .container { display: flex; gap: 30px; }
            .sidebar { width: 30%; background: #f9fafb; padding: 30px; border-radius: 20px; }
            .main { width: 70%; padding-top: 10px; }
            h2 { border-bottom: 2px solid var(--accent-color); font-size: 14px; background: none; }
            h1 { font-size: 24px; color: var(--accent-color); }
        </style></head>
        <body>
            <div class="container">
                <div class="sidebar">
                    <h1>{{ header.name }}</h1>
                    <p style="font-weight: bold; margin-bottom: 20px;">{{ header.title }}</p>
                    <div style="font-size: 11px; margin-bottom: 30px;">
                        <div>{{ header.email }}</div>
                        <div>{{ header.phone }}</div>
                        <div>{{ header.location }}</div>
                    </div>

                    {% if skills %}
                    <h2>Expertise</h2>
                    <ul style="list-style: none; padding:0; font-size: 11px;">
                        {% for s in skills %}<li style="margin-bottom: 4px;">• {{ s }}</li>{% endfor %}
                    </ul>
                    {% endif %}

                    {% if education %}
                    <h2>Education</h2>
                    <div style="font-size: 10px;">
                        {% for edu in education %}
                        <div style="margin-bottom: 10px;">
                            <b>{{ edu.school }}</b><br>{{ edu.degree }}
                        </div>
                        {% endfor %}
                    </div>
                    {% endif %}
                </div>
                <div class="main">
                    {% if summary %}<h2>Profile</h2><p style="font-size: 12px;">{{ summary }}</p>{% endif %}

                    {% if work_history %}
                    <h2>Experience</h2>
                    {% for job in work_history %}
                    <div class="section-item">
                        <div style="font-weight: bold;">{{ job.title }}</div>
                        <div style="font-size: 12px; color: #666;">{{ job.company }} | {{ job.start_date }} - {{ job.end_date }}</div>
                        <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
                    </div>
                    {% endfor %}
                    {% endif %}

                    {% if projects %}
                    <h2>Projects</h2>
                    {% for p in projects %}
                    <div class="section-item">
                        <div style="font-weight: bold;">{{ p.name }}</div>
                        <ul class="bullets">{% for b in p.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
                    </div>
                    {% endfor %}
                    {% endif %}
                </div>
            </div>
        </body>
        </html>
        """,
        "elegant": """
        <html>
        <head><style>
            {{ css }}
            body { font-family: 'Lato', sans-serif; }
            h1 { letter-spacing: 4px; text-align: center; border-bottom: 1px solid #eee; padding-bottom: 10px; }
            h2 { color: var(--accent-color); text-align: center; border-bottom: none; }
            h2::after { content: ''; display: block; width: 40px; height: 2px; background: var(--accent-color); margin: 5px auto; }
            .section-item { text-align: center; margin-bottom: 30px; }
            .bullets { list-style: none; padding: 0; }
        </style></head>
        <body>
            <h1>{{ header.name | upper }}</h1>
            <p style="text-align: center; font-size: 13px;">{{ header.email }} &nbsp;•&nbsp; {{ header.phone }} &nbsp;•&nbsp; {{ header.location }}</p>

            {% if summary %}<h2>Professional Summary</h2><p style="text-align: center; max-width: 600px; margin: 0 auto;">{{ summary }}</p>{% endif %}

            {% if work_history %}
            <h2>Experience</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div style="font-weight: bold; font-size: 1.1em;">{{ job.title }}</div>
                <div style="font-style: italic; color: #666;">{{ job.company }}, {{ job.location }} ({{ job.start_date }} - {{ job.end_date }})</div>
                <div style="margin-top: 10px;">
                    {% for b in job.bullets %}{% if b.strip() %}<div style="margin-bottom: 4px;">{{ b }}</div>{% endif %}{% endfor %}
                </div>
            </div>
            {% endfor %}
            {% endif %}

            {% if education %}
            <h2>Education</h2>
            {% for edu in education %}
            <div class="section-item">
                <b>{{ edu.school }}</b><br>{{ edu.degree }} ({{ edu.date }})
            </div>
            {% endfor %}
            {% endif %}

            {% if skills %}
            <h2>Key Skills</h2>
            <p style="text-align: center;">{{ skills | join(' | ') }}</p>
            {% endif %}
        </body>
        </html>
        """,
        "creative": """
        <html>
        <head><style>
            {{ css }}
            body { background: #1a1a1a; color: #ffffff; padding: 0; }
            .side-bar { background: var(--accent-color); width: 15px; position: fixed; left: 0; top: 0; bottom: 0; }
            .content { padding: 60px 80px; }
            h1 { font-size: 56px; line-height: 1; margin: 0; }
            h2 { color: var(--accent-color); border-bottom: 2px solid #333; font-size: 24px; margin-top: 40px; }
            .skill-tag { background: #333; border: none; color: white; }
            .item-header { color: var(--accent-color); }
        </style></head>
        <body>
            <div class="side-bar"></div>
            <div class="content">
                <h1>{{ header.name }}</h1>
                <p style="font-size: 24px; opacity: 0.6;">{{ header.title }}</p>
                <div style="margin-top: 20px; font-size: 14px; opacity: 0.8;">{{ header.email }} / {{ header.linkedin }} / {{ header.location }}</div>

                {% if summary %}
                <h2>Profile</h2>
                <p style="font-size: 16px; line-height: 1.8; opacity: 0.9;">{{ summary }}</p>
                {% endif %}

                {% if work_history %}
                <h2>Work Experience</h2>
                {% for job in work_history %}
                <div class="section-item" style="margin-bottom: 30px;">
                    <div class="item-header" style="font-size: 18px; font-weight: 800;">{{ job.company | upper }} // {{ job.title }}</div>
                    <div style="font-size: 12px; opacity: 0.5; margin-bottom: 10px;">{{ job.start_date }} - {{ job.end_date }} | {{ job.location }}</div>
                    <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
                </div>
                {% endfor %}
                {% endif %}

                {% if skills %}
                <h2>Skillset</h2>
                <div class="skills-list">{% for s in skills %}<span class="skill-tag">{{ s }}</span>{% endfor %}</div>
                {% endif %}

                {% if education %}
                <h2>Education</h2>
                {% for edu in education %}
                <div style="margin-bottom: 15px;">
                    <div style="font-weight: bold;">{{ edu.school }}</div>
                    <div style="opacity: 0.7;">{{ edu.degree }} ({{ edu.date }})</div>
                </div>
                {% endfor %}
                {% endif %}
            </div>
        </body>
        </html>
        """,
        "technical": """
        <html>
        <head><style>
            {{ css }}
            body { font-family: 'Inter', sans-serif; }
            code { background: #f3f4f6; padding: 2px 4px; border-radius: 4px; font-family: monospace; }
            h2 { color: #1e40af; border-bottom: 1px solid #d1d5db; font-size: 14px; text-transform: uppercase; margin-top: 25px; }
            .skill-tag { background: #eff6ff; color: #1e40af; border: 1px solid #bfdbfe; font-weight: 600; }
            .item-header { font-weight: 800; font-size: 1.1em; color: #111; }
        </style></head>
        <body>
            <div style="display: flex; justify-content: space-between; align-items: flex-start; border-bottom: 3px solid #111; padding-bottom: 15px;">
                <div><h1 style="margin:0; font-size: 28px;">{{ header.name }}</h1><p style="font-weight: 600; color: #444; font-size: 16px;">{{ header.title }}</p></div>
                <div style="text-align: right; font-size: 11px; line-height: 1.5; font-family: monospace;">
                    <div>{{ header.github }}</div>
                    <div>{{ header.linkedin }}</div>
                    <div>{{ header.email }}</div>
                    <div>{{ header.phone }}</div>
                </div>
            </div>

            {% if summary %}
            <h2>0x01. Summary</h2>
            <p style="font-size: 12px; color: #333;">{{ summary }}</p>
            {% endif %}

            {% if skills %}
            <h2>0x02. Technical Inventory</h2>
            <div class="skills-list">{% for s in skills %}<span class="skill-tag">{{ s }}</span>{% endfor %}</div>
            {% endif %}

            {% if work_history %}
            <h2>0x03. Professional History</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div class="item-header"><span>{{ job.company }}</span><span style="float: right; font-weight: normal; font-size: 0.8em;">{{ job.start_date }} - {{ job.end_date }}</span></div>
                <div style="font-weight: 600; color: #666; font-size: 0.9em; margin-bottom: 5px;">{{ job.title }}</div>
                <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if projects %}
            <h2>0x04. Key Projects</h2>
            {% for p in projects %}
            <div class="section-item">
                <div class="item-header"><span>{{ p.name }}</span><span style="float: right; font-weight: normal; font-size: 0.8em;">{{ p.date }}</span></div>
                <div class="item-sub">Role: {{ p.role }}</div>
                <ul class="bullets">{% for b in p.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if education %}
            <h2>0x05. Education</h2>
            {% for edu in education %}
            <div class="section-item" style="margin-bottom: 10px;">
                <div style="font-weight: 700;">{{ edu.school }}</div>
                <div style="font-size: 0.9em;">{{ edu.degree }} | {{ edu.date }}</div>
            </div>
            {% endfor %}
            {% endif %}
        </body>
        </html>
        """,
        "academic": """
        <html>
        <head><style>
            {{ css }}
            body { font-family: 'Times New Roman', serif; font-size: 14px; line-height: 1.4; color: black; padding: 40px 60px; }
            h1 { text-align: center; text-transform: uppercase; font-size: 22px; border: none; margin-bottom: 5px; }
            h2 { text-align: center; font-size: 16px; border-top: 1px solid black; border-bottom: 1px solid black; background: #f2f2f2; padding: 4px; margin-top: 30px; text-transform: uppercase; letter-spacing: 1px; }
            .section-item { margin-bottom: 20px; }
            .item-header { font-weight: bold; border-bottom: 0.5px solid #eee; margin-bottom: 5px; }
        </style></head>
        <body>
            <h1>{{ header.name }}</h1>
            <p style="text-align: center; font-size: 12px; margin-bottom: 30px;">
                {{ header.location }}<br>
                {{ header.email }} | {{ header.phone }}<br>
                {{ header.website }}
            </p>

            {% if summary %}
            <h2>Research Profile</h2>
            <p>{{ summary }}</p>
            {% endif %}

            {% if education %}
            <h2>Education</h2>
            {% for edu in education %}
            <div class="section-item">
                <div class="item-header"><span>{{ edu.school }}</span><span style="float: right;">{{ edu.date }}</span></div>
                <div style="font-style: italic;">{{ edu.degree }}</div>
                {% if edu.location %}<div style="font-size: 0.9em;">{{ edu.location }}</div>{% endif %}
            </div>
            {% endfor %}
            {% endif %}

            {% if work_history %}
            <h2>Professional Experience</h2>
            {% for job in work_history %}
            <div class="section-item">
                <div class="item-header"><span>{{ job.company }}</span><span style="float: right;">{{ job.start_date }} - {{ job.end_date }}</span></div>
                <div style="font-weight: bold; font-size: 0.9em;">{{ job.title }}</div>
                <ul class="bullets">{% for b in job.bullets %}{% if b.strip() %}<li>{{ b }}</li>{% endif %}{% endfor %}</ul>
            </div>
            {% endfor %}
            {% endif %}

            {% if publications %}
            <h2>Publications & Patents</h2>
            <div style="font-size: 13px;">
                {% for pub in publications %}
                <div style="margin-bottom: 12px; padding-left: 20px; text-indent: -20px;">{{ pub.authors or header.name }} ({{ pub.date }}). "{{ pub.title }}." <i>{{ pub.journal or pub.publisher }}</i>.</div>
                {% endfor %}
            </div>
            {% endif %}

            {% if skills %}
            <h2>Technical Skills</h2>
            <p>{{ skills | join(', ') }}</p>
            {% endif %}
        </body>
        </html>
        """,
        "cover_letter_standard": """
        <html>
        <head><style>
            {{ css }}
            .cl-body { margin-top: 40px; font-size: 14px; line-height: 1.6; }
            .cl-paragraph { margin-bottom: 24px; text-align: justify; }
            .cl-sign-off { margin-top: 40px; }
        </style></head>
        <body>
            <div style="text-align: center;">
                <h1 style="margin: 0; padding: 0;">{{ header.name }}</h1>
                <div class="header-sub">
                    {{ header.email }} | {{ header.phone }} | {{ header.location }}
                </div>
            </div>

            <div class="cl-body">
                <p class="cl-paragraph">{{ salutation }}</p>
                <p class="cl-paragraph">{{ opening }}</p>
                <p class="cl-paragraph">{{ why_us }}</p>
                <p class="cl-paragraph">{{ experience_highlight }}</p>
                <p class="cl-paragraph">{{ closing }}</p>
                <div class="cl-sign-off">
                    {{ sign_off }}<br>
                    <b style="font-size: 1.1em;">{{ header.name }}</b>
                </div>
            </div>
        </body>
        </html>
        """,
        "interview_session_standard": """
        <html>
        <head><style>
            {{ css }}
            .session-header { border-bottom: 4px solid var(--accent-color); padding-bottom: 20px; margin-bottom: 40px; }
            .question-card { background: #f9fafb; border: 1px solid #e5e7eb; border-radius: 12px; padding: 24px; margin-bottom: 24px; page-break-inside: avoid; }
            .score-badge { display: inline-block; background: var(--accent-color); color: white; padding: 4px 12px; border-radius: 99px; font-weight: bold; font-size: 12px; }
            .feedback-section { margin-top: 16px; font-size: 13px; color: #4b5563; }
        </style></head>
        <body>
            <div class="session-header">
                <h1>Interview Prep Summary: {{ name }}</h1>
                <div class="header-sub">Difficulty: {{ difficulty }} | Overall Score: {{ overall_score }}/10</div>
            </div>

            {% for q in questions %}
            <div class="question-card">
                <div style="display: flex; justify-content: space-between; align-items: center; margin-bottom: 12px;">
                    <span style="font-size: 10px; font-weight: 800; text-transform: uppercase; color: #6b7280;">{{ q.category }}</span>
                    <div class="score-badge">Score: {{ q.score }}/10</div>
                </div>
                <h3 style="margin-bottom: 16px;">{{ q.question_text }}</h3>
                <div style="font-style: italic; color: #6b7280; margin-bottom: 16px;">"{{ q.user_answer }}"</div>

                <div class="feedback-section">
                    <b>AI Feedback:</b><br>
                    {{ q.feedback.suggestions }}
                </div>

                {% if q.improved_answer %}
                <div class="feedback-section" style="background: #eff6ff; padding: 15px; border-radius: 8px; border-left: 4px solid #3b82f6; margin-top: 20px;">
                    <b style="color: #1e40af;">Recommended Answer:</b><br>
                    {{ q.improved_answer }}
                </div>
                {% endif %}
            </div>
            {% endfor %}
        </body>
        </html>
        """,
    }

    def render_to_html(
        self, profile: Dict[str, Any], template_id: str = "classic_ats", config: Dict[str, Any] = None
    ) -> str:
        """Renders profile data into an HTML string based on a template."""
        html_template = self.TEMPLATES.get(template_id, self.TEMPLATES["classic_ats"])

        # 1. Render BASE_CSS with config
        css_template = Template(self.BASE_CSS)
        rendered_css = css_template.render(config=config or {})

        # 2. Render final template
        template = Template(html_template)
        data = {
            "header": profile.get("header", {}),
            "summary": profile.get("summary", ""),
            "work_history": profile.get("work_history", []),
            "education": profile.get("education", []),
            "projects": profile.get("projects", []),
            "skills": profile.get("skills", []),
            "certifications": profile.get("certifications", []),
            "languages": profile.get("languages", []),
            "awards": profile.get("awards", []),
            "publications": profile.get("publications", []),
            "volunteer": profile.get("volunteer", []),
            "interests": profile.get("interests", []),
            "references": profile.get("references", []),
            "custom_sections": profile.get("custom_sections", []),
            "css": rendered_css,
            "config": config or {},
        }
        return template.render(**data)

    def render_cover_letter_to_html(
        self, content: Dict[str, Any], template_id: str = "cover_letter_standard", config: Dict[str, Any] = None
    ) -> str:
        """Renders cover letter content into a professional A4 HTML string."""
        html_template = self.TEMPLATES.get(template_id, self.TEMPLATES["cover_letter_standard"])
        css_template = Template(self.BASE_CSS)
        rendered_css = css_template.render(config=config or {})

        template = Template(html_template)
        data = {
            **content,
            "css": rendered_css,
        }
        return template.render(**data)

    def render_interview_session_to_html(
        self, session_data: Dict[str, Any], template_id: str = "interview_session_standard", config: Dict[str, Any] = None
    ) -> str:
        """Renders an interview prep session into a professional HTML report."""
        html_template = self.TEMPLATES.get(template_id, self.TEMPLATES["interview_session_standard"])
        css_template = Template(self.BASE_CSS)
        rendered_css = css_template.render(config=config or {})

        template = Template(html_template)
        data = {
            **session_data,
            "css": rendered_css,
        }
        return template.render(**data)

    def render_cover_letter_to_markdown(self, content: Dict[str, Any]) -> str:
        """Renders cover letter into a clean Markdown document."""
        header = content.get("header", {})
        md = f"# {header.get('name', 'Cover Letter')}\n"
        md += f"{header.get('email')} | {header.get('phone')} | {header.get('location')}\n\n"
        md += f"{content.get('salutation', 'Dear Hiring Manager,')}\n\n"
        md += f"{content.get('opening', '')}\n\n"
        md += f"{content.get('why_us', '')}\n\n"
        md += f"{content.get('experience_highlight', '')}\n\n"
        md += f"{content.get('closing', '')}\n\n"
        md += f"{content.get('sign_off', 'Best regards,')}\n\n"
        md += f"**{header.get('name')}**"
        return md

    async def export_pdf_cover_letter(self, content: Dict[str, Any], template_id: str, output_path: str):
        """Generates a PDF for a cover letter using playwright."""
        html = self.render_cover_letter_to_html(content, template_id)
        try:
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                page = await browser.new_page()
                await page.set_content(html)
                await page.pdf(path=output_path, format="A4", print_background=True)
                await browser.close()
            logger.info(f"CL PDF exported to {output_path}")
        except Exception as e:
            logger.error(f"CL PDF export failed: {e}")
            raise e

    async def export_pdf_interview_session(self, session_data: Dict[str, Any], template_id: str, output_path: str):
        """Generates a PDF for an interview session using playwright."""
        html = self.render_interview_session_to_html(session_data, template_id)
        try:
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                page = await browser.new_page()
                await page.set_content(html)
                await page.pdf(path=output_path, format="A4", print_background=True)
                await browser.close()
            logger.info(f"Interview PDF exported to {output_path}")
        except Exception as e:
            logger.error(f"Interview PDF export failed: {e}")
            raise e

    def render_to_markdown(self, profile: Dict[str, Any]) -> str:
        """Renders profile data into a clean Markdown string."""
        header = profile.get("header", {})
        md = f"# {header.get('name', 'Resume')}\n\n"
        md += f"**{header.get('title', '')}**\n\n"
        md += f"{header.get('email')} | {header.get('phone')} | {header.get('location')}\n"
        if header.get("linkedin"):
            md += f"[LinkedIn]({header.get('linkedin')}) "
        if header.get("github"):
            md += f"| [GitHub]({header.get('github')})"
        md += "\n\n---\n\n"

        if profile.get("summary"):
            md += f"## Professional Summary\n{profile.get('summary')}\n\n"

        if profile.get("work_history"):
            md += "## Experience\n"
            for job in profile.get("work_history", []):
                md += f"### {job.get('company')} | {job.get('title')}\n"
                md += f"*{job.get('start_date')} - {job.get('end_date')}* | {job.get('location')}\n\n"
                for b in job.get("bullets", []):
                    if b.strip():
                        md += f"- {b}\n"
                md += "\n"

        if profile.get("education"):
            md += "## Education\n"
            for edu in profile.get("education", []):
                md += f"### {edu.get('school')}\n"
                md += f"*{edu.get('degree')}* | {edu.get('date')} | {edu.get('location')}\n\n"

        if profile.get("projects"):
            md += "## Projects\n"
            for p in profile.get("projects", []):
                md += f"### {p.get('name')}\n"
                md += f"*{p.get('role')}* | {p.get('date')}\n\n"
                for b in p.get("bullets", []):
                    if b.strip():
                        md += f"- {b}\n"
                md += "\n"

        if profile.get("skills"):
            md += f"## Skills\n{', '.join(profile.get('skills', []))}\n\n"

        if profile.get("certifications"):
            md += "## Certifications\n"
            for c in profile.get("certifications", []):
                md += f"- **{c.get('name')}**, {c.get('issuer')} ({c.get('date')})\n"
            md += "\n"

        return md

    def export_docx(self, profile: Dict[str, Any], output_path: str):
        """Generates a high-quality Microsoft Word document."""
        if not Document:
            logger.error("python-docx not installed.")
            return

        doc = Document()

        # 1. Header
        header = profile.get("header", {})
        title = doc.add_heading(header.get("name", "Resume"), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{header.get('email')} | {header.get('phone')} | {header.get('location')}")

        # 2. Summary
        summary = profile.get("summary")
        if summary:
            doc.add_heading("Professional Summary", level=1)
            doc.add_paragraph(summary)

        # 3. Experience
        work = profile.get("work_history", [])
        if work:
            doc.add_heading("Experience", level=1)
            for job in work:
                p = doc.add_paragraph()
                p.add_run(f"{job.get('company')}").bold = True
                p.add_run(f"\t\t{job.get('start_date')} - {job.get('end_date')}").italic = True

                p2 = doc.add_paragraph()
                p2.add_run(f"{job.get('title')}").italic = True

                for bullet in job.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

        # 4. Education
        edu_list = profile.get("education", [])
        if edu_list:
            doc.add_heading("Education", level=1)
            for edu in edu_list:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('school')}").bold = True
                p.add_run(f"\t\t{edu.get('date')}").italic = True
                doc.add_paragraph(f"{edu.get('degree')} in {edu.get('location')}")

        # 5. Projects
        projects = profile.get("projects", [])
        if projects:
            doc.add_heading("Projects", level=1)
            for p in projects:
                p_item = doc.add_paragraph()
                p_item.add_run(f"{p.get('name')}").bold = True
                p_item.add_run(f"\t\t{p.get('date')}").italic = True
                for bullet in p.get("bullets", []):
                    doc.add_paragraph(bullet, style="List Bullet")

        # 6. Skills
        skills = profile.get("skills", [])
        if skills:
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(", ".join(skills))

        # 7. Certifications
        certs = profile.get("certifications", [])
        if certs:
            doc.add_heading("Certifications", level=1)
            for c in certs:
                doc.add_paragraph(f"{c.get('name')} - {c.get('issuer')} ({c.get('date')})")

        doc.save(output_path)
        logger.info(f"DOCX exported to {output_path}")

    async def export_pdf(
        self,
        profile: Dict[str, Any],
        template_id: str,
        output_path: str,
        config: Dict[str, Any] = None,
    ):
        """Generates a PDF using playwright (browser-based rendering)."""
        html_content = self.render_to_html(profile, template_id, config=config)

        try:
            from playwright.async_api import async_playwright
            async with async_playwright() as p:
                browser = await p.chromium.launch()
                page = await browser.new_page()
                await page.set_content(html_content)
                await page.pdf(path=output_path, format="A4", print_background=True)
                await browser.close()
            logger.info(f"PDF exported to {output_path}")
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            raise e


template_engine = TemplateEngine()
